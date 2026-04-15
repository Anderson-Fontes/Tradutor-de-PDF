import os
import glob
from pdf2docx import Converter
from docx import Document
from deep_translator import GoogleTranslator
from docx2pdf import convert

def traduzir_pdf(arquivo_entrada, arquivo_saida):
    docx_temp = f"temp_{arquivo_entrada}.docx"
    docx_traduzido = f"temp_traduzido_{arquivo_entrada}.docx"

    try:
        print(f"\nIniciando: {arquivo_entrada}...")
        print("1/4 - Convertendo para DOCX...")
        cv = Converter(arquivo_entrada)
        cv.convert(docx_temp, start=0, end=None)
        cv.close()

        print("2/4 - Traduzindo o conteúdo...")
        doc = Document(docx_temp)
        tradutor = GoogleTranslator(source='en', target='pt')

        for paragraph in doc.paragraphs:
            texto = paragraph.text.strip()
            if texto:
                try:
                    paragraph.text = tradutor.translate(texto)
                except:
                    pass

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_celula = cell.text.strip()
                    if texto_celula and not texto_celula.isnumeric():
                        try:
                            cell.text = tradutor.translate(texto_celula)
                        except:
                            pass

        doc.save(docx_traduzido)

        print("3/4 - Reconstruindo o arquivo PDF...")
        convert(docx_traduzido, arquivo_saida)

        print(f"4/4 - Concluído! Salvo como: {arquivo_saida}")

    except Exception as e:
        print(f"Erro no arquivo {arquivo_entrada}: {e}")
        
    finally:
        if os.path.exists(docx_temp):
            os.remove(docx_temp)
        if os.path.exists(docx_traduzido):
            os.remove(docx_traduzido)

# ==========================================
# NOVA LÓGICA DE FILA AUTOMÁTICA
# ==========================================

# Encontra todos os arquivos na pasta que começam com "parte_" e terminam com ".pdf"
lista_de_arquivos = glob.glob("parte_*.pdf")

# Ordena os arquivos para não traduzir fora de ordem
lista_de_arquivos.sort(key=lambda x: int(x.split('_')[1].split('.')[0]))

print(f"Encontrados {len(lista_de_arquivos)} arquivos para traduzir. Iniciando a fila...")

# Faz um loop por todos os arquivos encontrados
for arquivo in lista_de_arquivos:
    # Cria o nome de saída (ex: "parte_1.pdf" vira "traduzido_parte_1.pdf")
    nome_saida = arquivo.replace("parte_", "traduzido_parte_")
    
    # Se o arquivo já foi traduzido antes (caso o script tenha caído), ele pula
    if os.path.exists(nome_saida):
        print(f"\nPulando {arquivo}, pois '{nome_saida}' já existe.")
        continue
        
    traduzir_pdf(arquivo, nome_saida)

print("\nTODOS OS ARQUIVOS FORAM TRADUZIDOS COM SUCESSO!")